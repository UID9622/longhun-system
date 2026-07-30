#!/usr/bin/env python3
#龍芯⚡️丙午·丙申·癸酉·庚申·临-MD_TO_DOCX-v1.0-8950c909
# CREATOR: 诸葛鑫 (UID9622)
# PROTOCOL: CC BY-NC-SA 4.0
# -*- coding: utf-8 -*-
"""
龍魂·Markdown转Word (python-docx)
依赖: venv 内 python-docx + markdown
功能: 将带DNA追溯码的Markdown转为带龍魂样式的 .docx。

用法:
    source .venv_docs/bin/activate
    python3 bin/docs_factory/md_to_docx.py input.md output.docx
"""
import sys
import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DNA = "#CONFIRM🌌9622-ONLY-ONCE🧬LK9X-772Z"

def build(md_file, out):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Microsoft YaHei"
    st.font.size = Pt(12)

    with open(md_file, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    for ln in lines:
        if ln.startswith("# "):
            h = doc.add_heading(ln[2:], level=1)
            _dragonize(h)
        elif ln.startswith("## "):
            h = doc.add_heading(ln[3:], level=2)
            _dragonize(h)
        elif ln.startswith("### "):
            doc.add_heading(ln[4:], level=3)
        elif ln.strip() == "":
            continue
        else:
            doc.add_paragraph(ln)

    # DNA尾注
    dna = doc.add_paragraph()
    dna.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dna.add_run("DNA追溯码: " + DNA)
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(128, 128, 128)
    doc.save(out)
    return f"Word已生成: {out}"

def _dragonize(heading):
    for r in heading.runs:
        r.font.color.rgb = RGBColor(196, 30, 58)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("md"); ap.add_argument("out")
    a = ap.parse_args()
    print(build(a.md, a.out))

if __name__ == "__main__":
    main()
